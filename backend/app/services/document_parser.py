from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentParseError(ValueError):
    pass


@dataclass(slots=True)
class ParsedSection:
    text: str
    page: int | None = None
    section_title: str | None = None


class DocumentParser:
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx"}

    def parse(self, path: Path) -> list[ParsedSection]:
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise DocumentParseError(f"不支持的文件类型：{extension}")
        parser = getattr(self, f"_parse_{extension[1:]}")
        sections = [item for item in parser(path) if item.text.strip()]
        if not sections:
            if extension == ".pdf":
                raise DocumentParseError("PDF 未解析出文本，可能是扫描件；可后续接入 DeepSeek-OCR。")
            raise DocumentParseError("文档中没有可索引的文本")
        return sections

    @staticmethod
    def _read_text(path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("无法识别文本编码")

    def _parse_txt(self, path: Path) -> list[ParsedSection]:
        return [ParsedSection(self._read_text(path))]

    def _parse_md(self, path: Path) -> list[ParsedSection]:
        text = self._read_text(path)
        sections: list[ParsedSection] = []
        title: str | None = None
        buffer: list[str] = []
        for line in text.splitlines():
            heading = re.match(r"^#{1,6}\s+(.+)$", line.strip())
            if heading:
                if buffer:
                    sections.append(ParsedSection("\n".join(buffer), section_title=title))
                title, buffer = heading.group(1).strip(), []
            else:
                buffer.append(line)
        if buffer:
            sections.append(ParsedSection("\n".join(buffer), section_title=title))
        return sections

    @staticmethod
    def _parse_pdf(path: Path) -> list[ParsedSection]:
        reader = PdfReader(str(path))
        return [ParsedSection(page.extract_text() or "", page=index) for index, page in enumerate(reader.pages, 1)]

    @staticmethod
    def _parse_docx(path: Path) -> list[ParsedSection]:
        document = DocxDocument(str(path))
        sections: list[ParsedSection] = []
        title: str | None = None
        buffer: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                if buffer:
                    sections.append(ParsedSection("\n".join(buffer), section_title=title))
                title, buffer = text, []
            else:
                buffer.append(text)
        if buffer:
            sections.append(ParsedSection("\n".join(buffer), section_title=title))
        for table_index, table in enumerate(document.tables, 1):
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            sections.append(ParsedSection("\n".join(rows), section_title=f"表格 {table_index}"))
        return sections

    @staticmethod
    def _frame_to_sections(frame: pd.DataFrame, title: str) -> list[ParsedSection]:
        frame = frame.fillna("")
        headers = [str(column) for column in frame.columns]
        lines = [" | ".join(f"{header}: {value}" for header, value in zip(headers, row)) for row in frame.astype(str).values]
        return [ParsedSection("\n".join(lines), section_title=title)]

    def _parse_csv(self, path: Path) -> list[ParsedSection]:
        last_error: Exception | None = None
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return self._frame_to_sections(pd.read_csv(path, encoding=encoding), path.stem)
            except UnicodeDecodeError as exc:
                last_error = exc
        raise DocumentParseError(f"CSV 解析失败：{last_error}")

    def _parse_xlsx(self, path: Path) -> list[ParsedSection]:
        workbook = pd.read_excel(path, sheet_name=None)
        sections: list[ParsedSection] = []
        for sheet_name, frame in workbook.items():
            sections.extend(self._frame_to_sections(frame, f"工作表：{sheet_name}"))
        return sections
